#!/usr/bin/env python3
"""
Simple converter: HTML (saved-as-.doc) -> .docx using python-docx and BeautifulSoup.
Usage: python convert_doc_html_to_docx.py input.doc output.docx
If output omitted, saves as input.docx next to input.
"""
import sys
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document


def html_to_docx(input_path: Path, output_path: Path):
    html = input_path.read_text(encoding='utf-8', errors='ignore')
    soup = BeautifulSoup(html, 'lxml')
    doc = Document()

    # If the HTML has a head/title, add as title
    title_tag = soup.find('title')
    if title_tag and title_tag.text.strip():
        doc.add_heading(title_tag.text.strip(), level=1)

    # Find the body; if none, use whole document
    body = soup.find('body') or soup

    for elem in body.children:
        if elem.name is None:
            # could be whitespace/newline
            text = str(elem).strip()
            if text:
                doc.add_paragraph(text)
            continue

        name = elem.name.lower()
        if name in ('h1', 'h2', 'h3'):
            level = 1 if name == 'h1' else (2 if name == 'h2' else 3)
            doc.add_heading(elem.get_text(strip=True), level=level)
        elif name == 'p':
            doc.add_paragraph(elem.get_text(separator=' ', strip=True))
        elif name in ('div', 'section'):
            # flatten basic div/section children
            text = elem.get_text(separator=' ', strip=True)
            if text:
                doc.add_paragraph(text)
        elif name == 'table':
            # convert table
            rows = elem.find_all('tr')
            if not rows:
                continue
            num_cols = max(len(r.find_all(['td','th'])) for r in rows)
            table = doc.add_table(rows=0, cols=num_cols)
            table.style = 'Table Grid'
            for r in rows:
                cells = r.find_all(['td','th'])
                row_cells = table.add_row().cells
                for i, c in enumerate(cells):
                    row_cells[i].text = c.get_text(separator=' ', strip=True)
        elif name in ('ul', 'ol'):
            for li in elem.find_all('li'):
                doc.add_paragraph(li.get_text(separator=' ', strip=True), style='List Bullet' if name=='ul' else 'List Number')
        elif name == 'br':
            doc.add_paragraph('')
        else:
            # fallback: append text
            text = elem.get_text(separator=' ', strip=True)
            if text:
                doc.add_paragraph(text)

    doc.save(str(output_path))


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python convert_doc_html_to_docx.py input.doc [output.docx]')
        sys.exit(2)
    input_path = Path(sys.argv[1])
    if not input_path.exists():
        print('Input file not found:', input_path)
        sys.exit(3)
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = input_path.with_suffix('.docx')
    try:
        html_to_docx(input_path, output_path)
        print('Saved', output_path)
    except Exception as e:
        print('Conversion failed:', e)
        sys.exit(4)
